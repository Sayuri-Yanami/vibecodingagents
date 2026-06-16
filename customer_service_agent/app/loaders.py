from __future__ import annotations

import csv
import json
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable


TEXT_EXTENSIONS = {".txt", ".md", ".markdown"}
SUPPORTED_EXTENSIONS = TEXT_EXTENSIONS | {".csv", ".json", ".pdf", ".docx"}


@dataclass(frozen=True)
class LoadedDocument:
    source: str
    text: str


def _read_text(path: Path) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError(f"无法识别文件编码：{path}")


def _read_csv(path: Path) -> str:
    content = _read_text(path)
    rows = []
    for row in csv.reader(content.splitlines()):
        cells = [cell.strip() for cell in row if cell.strip()]
        if cells:
            rows.append(" | ".join(cells))
    return "\n".join(rows)


def _read_json(path: Path) -> str:
    data = json.loads(_read_text(path))
    return json.dumps(data, ensure_ascii=False, indent=2)


def _read_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError(
            f"读取 PDF 需要可选依赖 pypdf：python -m pip install pypdf（文件：{path.name}）"
        ) from exc
    reader = PdfReader(str(path))
    return "\n\n".join(page.extract_text() or "" for page in reader.pages)


def _read_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(
            f"读取 DOCX 需要可选依赖 python-docx："
            f"python -m pip install python-docx（文件：{path.name}）"
        ) from exc
    document = Document(str(path))
    return "\n".join(paragraph.text for paragraph in document.paragraphs)


def read_document(path: Path, root: Path) -> LoadedDocument:
    suffix = path.suffix.lower()
    if suffix in TEXT_EXTENSIONS:
        text = _read_text(path)
    elif suffix == ".csv":
        text = _read_csv(path)
    elif suffix == ".json":
        text = _read_json(path)
    elif suffix == ".pdf":
        text = _read_pdf(path)
    elif suffix == ".docx":
        text = _read_docx(path)
    else:
        raise ValueError(f"不支持的文件格式：{path.name}")

    return LoadedDocument(
        source=path.relative_to(root).as_posix(),
        text=text.strip(),
    )


def iter_document_paths(root: Path) -> Iterable[Path]:
    if not root.exists():
        return []
    return (
        path
        for path in sorted(root.rglob("*"))
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS
    )


def load_documents(root: Path) -> list[LoadedDocument]:
    documents = []
    errors = []
    for path in iter_document_paths(root):
        try:
            document = read_document(path, root)
            if document.text:
                documents.append(document)
        except Exception as exc:
            errors.append(str(exc))

    if errors:
        joined = "\n- ".join(errors)
        raise RuntimeError(f"部分文档读取失败：\n- {joined}")
    return documents
